#!/usr/bin/env python3
"""
Gera o E-book 2 reescrito: "Novas Regras Cripto 2026"
Versão corrigida com base na auditoria profissional.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

OUTPUT_PATH = "/Users/user/Downloads/EBOOK_2_NOVAS_REGRAS_CRIPTO_2026.docx"

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x0B, 0x2A, 0x5C)
    if level == 1:
        hs.font.size = Pt(22)
        hs.font.bold = True
    elif level == 2:
        hs.font.size = Pt(16)
        hs.font.bold = True
    elif level == 3:
        hs.font.size = Pt(13)
        hs.font.bold = True
    elif level == 4:
        hs.font.size = Pt(11)
        hs.font.bold = True

# Helper to add styled paragraphs
def add_p(text, bold=False, italic=False, size=None, color=None, align=None, space_after=None, style_name=None):
    if style_name:
        p = doc.add_paragraph(style=style_name)
    else:
        p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_boost_opinion(text):
    """Marca opinião da Boost Research claramente."""
    p = doc.add_paragraph()
    run = p.add_run("📌 Visão Boost Research: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x0B, 0x6E, 0x4F)
    run2 = p.add_run(text)
    run2.italic = True
    run2.font.color.rgb = RGBColor(0x0B, 0x6E, 0x4F)
    return p

def add_legal_box(text):
    """Caixa de referência legal."""
    p = doc.add_paragraph()
    run = p.add_run(f"⚖️ Base legal: {text}")
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    return p

def add_numbered(text):
    p = doc.add_paragraph(text, style='List Number')
    return p


# ═══════════════════════════════════════════════════════════════════════
# CAPA
# ═══════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
add_p("BOOST RESEARCH", bold=True, size=14, color=(0x0B, 0x6E, 0x4F), align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_p("Novas Regras Cripto 2026", bold=True, size=28, color=(0x0B, 0x2A, 0x5C), align=WD_ALIGN_PARAGRAPH.CENTER)
add_p("O que o Investidor com Patrimônio Precisa\nSaber Antes de Julho", bold=True, size=16, color=(0x33, 0x33, 0x33), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
add_p("Guia completo com base na legislação vigente — atualizado em abril de 2026", italic=True, size=12, color=(0x66, 0x66, 0x66), align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_p("André Franco", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_p("Fundador e analista, Boost Research", italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
doc.add_paragraph()
add_p("© 2026 Boost Research. Todos os direitos reservados.", size=9, color=(0x99, 0x99, 0x99), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# DISCLAIMER
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("Aviso Legal", level=2)

disclaimer_text = (
    "Este material tem caráter exclusivamente educacional e informativo. "
    "Não constitui aconselhamento jurídico, tributário ou financeiro individualizado. "
    "As informações aqui apresentadas baseiam-se na legislação vigente até abril de 2026 "
    "e podem sofrer alterações em decorrência de novas normas, regulamentações ou "
    "interpretações dos órgãos competentes.\n\n"
    "Para decisões específicas sobre sua situação patrimonial, consulte um advogado "
    "tributarista e/ou um contador especializado em ativos digitais.\n\n"
    "A Boost Research é uma empresa de análise de criptoativos. As opiniões e recomendações "
    "comerciais expressas neste material estão claramente identificadas como \"Visão Boost Research\" "
    "e refletem a opinião dos analistas da empresa, não devendo ser confundidas com orientação "
    "legal ou tributária.\n\n"
    "Criptoativos são investimentos de alto risco. Resultados passados não garantem resultados futuros."
)
add_p(disclaimer_text, size=10, color=(0x66, 0x66, 0x66))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SUMÁRIO
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("Sumário", level=1)
capitulos = [
    "Capítulo 1 — O Que Mudou e Por Que Importa",
    "Capítulo 2 — DeCripto: A Nova Declaração Obrigatória",
    "Capítulo 3 — CARF/OCDE: Coleta vs Troca Internacional",
    "Capítulo 4 — O Novo Perímetro Regulatório do Banco Central",
    "Capítulo 5 — Tributação de Criptoativos: Guia com Árvore Decisória",
    "Capítulo 6 — Rearp: O Que Foi e O Que Significa Agora",
    "Capítulo 7 — Cenários de Multa e Penalidades",
    "Capítulo 8 — Checklist: 7 Passos para Conformidade",
    "Capítulo 9 — Próximos Passos",
    "Referências e Fontes",
]
for c in capitulos:
    add_p(c, size=12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# CAPÍTULO 1
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("Capítulo 1", level=1)
doc.add_heading("O Que Mudou e Por Que Importa", level=2)

add_p(
    "O mercado de criptoativos no Brasil passou por uma transformação regulatória sem precedentes "
    "entre 2022 e 2026. O que antes era um ambiente com poucas regras claras tornou-se um dos "
    "ecossistemas mais regulados da América Latina. Este capítulo apresenta a linha do tempo "
    "completa das mudanças e explica por que cada uma delas importa para o investidor com "
    "patrimônio relevante."
)

doc.add_heading("1.1 Linha do Tempo Regulatória", level=3)

timeline = [
    ("Dezembro de 2022 — Lei 14.478/2022 (Marco Legal dos Ativos Virtuais)",
     "Estabeleceu as diretrizes gerais para prestação de serviços de ativos virtuais no Brasil. "
     "Nos termos do art. 2º da Lei 14.478/2022, definiu-se ativo virtual como a representação "
     "digital de valor que pode ser negociada ou transferida por meios eletrônicos. "
     "A lei atribuiu ao Poder Executivo a competência para designar o órgão regulador."),

    ("Novembro de 2023 — Decreto 11.563/2023",
     "De acordo com o art. 1º do Decreto 11.563/2023, o Banco Central do Brasil foi designado "
     "como órgão responsável por regular, autorizar e supervisionar as prestadoras de serviços "
     "de ativos virtuais (PSAVs). Importante: a CVM somente atua como reguladora quando o "
     "ativo virtual se enquadra como valor mobiliário, nos termos do art. 2º da Lei 6.385/1976."),

    ("Dezembro de 2023 — Lei 14.754/2023",
     "Alterou a tributação de aplicações financeiras no exterior por pessoas físicas residentes "
     "no Brasil. Nos termos da Lei 14.754/2023, criptoativos mantidos em custódia no exterior "
     "passaram a ser tributados a 15% sobre os rendimentos, independentemente do valor de alienação."),

    ("Janeiro de 2025 — IN RFB 2.291/2025 (Cria a DeCripto)",
     "A Instrução Normativa RFB nº 2.291, de 22 de janeiro de 2025, instituiu a Declaração "
     "de Criptoativos (DeCripto), que substituirá a antiga obrigação acessória prevista na "
     "IN RFB 1.888/2019. A DeCripto entrará em vigor em 1º de julho de 2026."),

    ("Novembro de 2025 — Resoluções BCB 519, 520 e 521",
     "O Banco Central publicou três resoluções estruturantes para o setor:\n"
     "• Resolução BCB nº 519: estabelece requisitos para autorização e funcionamento das PSAVs.\n"
     "• Resolução BCB nº 520: define regras de governança, gestão de riscos e controles internos.\n"
     "• Resolução BCB nº 521: disciplina a segregação patrimonial entre recursos dos clientes e da PSAV."),

    ("Fevereiro de 2026 — Resolução CMN 5.280",
     "Nos termos da Resolução CMN nº 5.280, de fevereiro de 2026, as PSAVs autorizadas pelo "
     "Banco Central passaram a ser equiparadas a instituições financeiras para fins da Lei "
     "Complementar nº 105/2001 (sigilo bancário). Isso significa que o Banco Central e a "
     "Receita Federal podem acessar dados de clientes de PSAVs nas mesmas condições aplicáveis "
     "a bancos e corretoras tradicionais."),

    ("Julho de 2026 — DeCripto entra em vigor",
     "De acordo com a IN RFB 2.291/2025, a partir de 1º de julho de 2026 as exchanges e "
     "prestadoras de serviços de ativos virtuais passarão a reportar dados detalhados das "
     "operações de seus clientes à Receita Federal do Brasil por meio da DeCripto. "
     "Esta data marca o início da COLETA de dados, não da troca internacional."),

    ("2027 — Primeiras trocas internacionais via CARF",
     "As primeiras trocas automáticas de informações entre países sob o Crypto-Asset Reporting "
     "Framework (CARF) da OCDE estão previstas para 2027. É fundamental não confundir: julho "
     "de 2026 é o início da coleta doméstica via DeCripto; 2027 marca o início das trocas "
     "internacionais entre jurisdições que aderiram ao CARF.")
]

for title, desc in timeline:
    doc.add_heading(title, level=4)
    add_p(desc)

doc.add_heading("1.2 Por Que Isso Importa Para Você", level=3)

add_p(
    "Cada uma dessas normas amplia o perímetro de informações que o Estado brasileiro possui "
    "sobre operações com criptoativos. Para o investidor com patrimônio relevante, as implicações "
    "são concretas:"
)

implications = [
    "Transparência total: a partir de julho de 2026, toda operação realizada em exchanges brasileiras "
    "será reportada automaticamente à Receita Federal, nos termos da IN RFB 2.291/2025.",
    "Fiscalização cruzada: com as trocas internacionais via CARF previstas para 2027, operações "
    "em exchanges estrangeiras também serão visíveis ao Fisco brasileiro.",
    "Equiparação bancária: nos termos da Resolução CMN 5.280, PSAVs são agora equiparadas a "
    "instituições financeiras para fins de sigilo, permitindo acesso regulatório aos dados.",
    "Tributação de ativos no exterior: a Lei 14.754/2023 eliminou a possibilidade de diferir "
    "tributação de criptoativos mantidos fora do Brasil.",
    "Penalidades relevantes: multas por omissão ou incorreção de dados na DeCripto podem chegar "
    "a 3% do valor da operação, conforme previsto na IN RFB 2.291/2025."
]
for imp in implications:
    add_bullet(imp)

add_boost_opinion(
    "O investidor que se antecipa à vigência da DeCripto em julho de 2026, organizando seu "
    "portfólio e histórico de operações, estará em posição significativamente melhor do que "
    "aquele que espera a obrigação se tornar efetiva para agir."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# CAPÍTULO 2
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("Capítulo 2", level=1)
doc.add_heading("DeCripto: A Nova Declaração Obrigatória", level=2)

add_p(
    "A Declaração de Criptoativos (DeCripto) é a principal novidade operacional para o mercado "
    "cripto brasileiro em 2026. Instituída pela Instrução Normativa RFB nº 2.291, de 22 de "
    "janeiro de 2025, ela representa um salto qualitativo na capacidade de fiscalização da "
    "Receita Federal sobre operações com ativos virtuais."
)

doc.add_heading("2.1 O Que É a DeCripto", level=3)
add_p(
    "A DeCripto é uma obrigação acessória que substitui o modelo anterior de reporte previsto "
    "na IN RFB 1.888/2019. Nos termos da IN RFB 2.291/2025, a DeCripto amplia "
    "significativamente o escopo e o detalhamento das informações reportadas."
)
add_legal_box("IN RFB nº 2.291, de 22 de janeiro de 2025 — institui a DeCripto.")

doc.add_heading("2.2 Quem Deve Reportar", level=3)
add_p("De acordo com a IN RFB 2.291/2025, são obrigados a prestar informações via DeCripto:")
obrigados = [
    "Exchanges e prestadoras de serviços de ativos virtuais (PSAVs) autorizadas a operar no Brasil, "
    "nos termos da Resolução BCB nº 519/2025.",
    "Pessoas físicas residentes no Brasil que realizem operações em exchanges não domiciliadas "
    "(estrangeiras), quando o valor mensal das operações ultrapassar o limite estabelecido pela "
    "Receita Federal.",
    "Pessoas jurídicas brasileiras que realizem operações com criptoativos, diretamente ou por "
    "meio de custodiantes no exterior."
]
for o in obrigados:
    add_bullet(o)

doc.add_heading("2.3 Quais Dados São Reportados", level=3)
add_p("A DeCripto exige, nos termos da IN RFB 2.291/2025, o reporte das seguintes informações:")
dados = [
    "Identificação completa do titular (CPF/CNPJ, nome, endereço).",
    "Tipo de criptoativo negociado (com identificação específica — ex.: BTC, ETH, stablecoins).",
    "Tipo de operação (compra, venda, permuta, transferência, staking, entre outros).",
    "Data e valor da operação em reais.",
    "Quantidade de criptoativo transacionado.",
    "Identificação da contraparte, quando disponível.",
    "Saldos mantidos em custódia ao final de cada período.",
    "Endereços de carteira utilizados nas operações (wallet addresses)."
]
for d in dados:
    add_bullet(d)

add_p(
    "O nível de detalhamento é significativamente superior ao exigido pela IN RFB 1.888/2019, "
    "que se limitava a informações agregadas sobre operações de compra e venda."
)

doc.add_heading("2.4 Quando Começa", level=3)
add_p(
    "De acordo com a IN RFB 2.291/2025, a DeCripto entrará em vigor em 1º de julho de 2026. "
    "A partir dessa data, as informações deverão ser prestadas mensalmente pelas entidades "
    "obrigadas."
)
add_p(
    "Ponto crucial: julho de 2026 marca o início da coleta doméstica de dados pela Receita "
    "Federal do Brasil. As trocas internacionais de informações sob o CARF da OCDE estão "
    "previstas para 2027 (ver Capítulo 3).",
    bold=True
)

doc.add_heading("2.5 Penalidades por Descumprimento", level=3)
add_p("Nos termos da IN RFB 2.291/2025, o descumprimento da DeCripto pode acarretar:")
penalidades_decripto = [
    "Multa por atraso na entrega: valores fixos progressivos conforme o tempo de atraso.",
    "Multa por informação incorreta ou omissa: até 3% do valor da operação omitida ou informada "
    "incorretamente.",
    "Multa por não atendimento a intimação: aplicável quando a Receita Federal solicitar "
    "esclarecimentos adicionais e a entidade obrigada não responder no prazo."
]
for pen in penalidades_decripto:
    add_bullet(pen)

add_legal_box(
    "IN RFB nº 2.291/2025 — disposições sobre penalidades. "
    "Ver também: art. 57 da Medida Provisória nº 2.158-35/2001 (penalidades gerais por "
    "descumprimento de obrigações acessórias)."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# CAPÍTULO 3
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("Capítulo 3", level=1)
doc.add_heading("CARF/OCDE: Coleta vs Troca Internacional", level=2)

add_p(
    "Um dos erros mais comuns na interpretação do cenário regulatório de 2026 é confundir "
    "dois eventos distintos: o início da coleta de dados via DeCripto (julho de 2026) e o "
    "início das trocas internacionais de informações via CARF (2027). Este capítulo esclarece "
    "essa distinção fundamental."
)

doc.add_heading("3.1 O Que É o CARF da OCDE", level=3)
add_p(
    "O Crypto-Asset Reporting Framework (CARF) é um padrão desenvolvido pela Organização "
    "para a Cooperação e Desenvolvimento Econômico (OCDE) para a troca automática de "
    "informações sobre criptoativos entre jurisdições. O CARF foi aprovado pelo G20 em 2023 "
    "e o Brasil está entre os países que se comprometeram com sua implementação."
)
add_p(
    "O CARF funciona de forma análoga ao Common Reporting Standard (CRS), já utilizado para "
    "troca de informações sobre contas financeiras tradicionais. A diferença é que o CARF "
    "foi desenhado especificamente para criptoativos, cobrindo exchanges, brokers e outros "
    "intermediários de ativos virtuais."
)

doc.add_heading("3.2 A Distinção Essencial: Coleta ≠ Troca", level=3)

add_p("É imprescindível separar dois marcos temporais distintos:", bold=True)

doc.add_heading("Julho de 2026 — Início da Coleta Doméstica", level=4)
add_p(
    "De acordo com a IN RFB 2.291/2025, a partir de 1º de julho de 2026 as exchanges e PSAVs "
    "brasileiras passarão a coletar e reportar dados à Receita Federal do Brasil por meio da "
    "DeCripto. Este é um processo doméstico: os dados são coletados por entidades brasileiras "
    "e enviados à Receita Federal."
)

doc.add_heading("2027 — Primeiras Trocas Internacionais", level=4)
add_p(
    "As primeiras trocas automáticas de informações entre países sob o CARF estão previstas "
    "para 2027. Nesse momento, dados coletados por jurisdições estrangeiras que aderiram ao "
    "CARF serão compartilhados com o Brasil — e vice-versa. Isso significa que operações "
    "realizadas por residentes brasileiros em exchanges domiciliadas em outros países signatários "
    "serão comunicadas à Receita Federal."
)

doc.add_heading("3.3 Implicações Práticas", level=3)

add_p("A separação entre coleta e troca tem consequências práticas relevantes:")
impl_carf = [
    "A partir de julho de 2026, suas operações em exchanges brasileiras serão reportadas "
    "automaticamente — independentemente da troca internacional.",
    "A partir de 2027, operações em exchanges estrangeiras de países que aderiram ao CARF "
    "também poderão ser comunicadas ao Fisco brasileiro.",
    "A combinação DeCripto + CARF cria um cerco informacional: operações domésticas e "
    "internacionais ficarão, progressivamente, visíveis ao regulador.",
    "Países que ainda não aderiram ao CARF podem fazê-lo a qualquer momento — a lista de "
    "jurisdições signatárias tende a crescer."
]
for i in impl_carf:
    add_bullet(i)

doc.add_heading("3.4 Países Comprometidos com o CARF", level=3)
add_p(
    "Até abril de 2026, mais de 50 jurisdições manifestaram compromisso com a implementação "
    "do CARF, incluindo todos os membros do G20. A lista completa e atualizada está disponível "
    "no site da OCDE (oecd.org/tax/automatic-exchange)."
)

add_boost_opinion(
    "O investidor que opera em múltiplas jurisdições deve considerar que, mesmo que uma exchange "
    "estrangeira não reporte dados em 2026, as trocas via CARF em 2027 poderão retroagir ao "
    "período de coleta. Organizar o histórico de operações agora é mais prudente do que aguardar."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# CAPÍTULO 4
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("Capítulo 4", level=1)
doc.add_heading("O Novo Perímetro Regulatório do Banco Central", level=2)

add_p(
    "Um dos pontos que mais gera confusão entre investidores e até profissionais do mercado "
    "é a questão de quem regula as prestadoras de serviços de ativos virtuais (PSAVs) no "
    "Brasil. Este capítulo esclarece definitivamente o papel do Banco Central e da CVM."
)

doc.add_heading("4.1 Banco Central: O Regulador Padrão das PSAVs", level=3)
add_p(
    "Nos termos do Decreto 11.563/2023, o Banco Central do Brasil foi designado como o "
    "órgão responsável por autorizar e supervisionar as PSAVs. Isso significa que toda empresa "
    "que presta serviços de intermediação, custódia ou negociação de criptoativos no Brasil "
    "deve, por padrão, obter autorização junto ao Banco Central."
)
add_legal_box(
    "Decreto nº 11.563, de 13 de junho de 2023 — designa o Banco Central como regulador "
    "das prestadoras de serviços de ativos virtuais."
)

doc.add_heading("4.2 Resoluções BCB 519, 520 e 521 (Novembro de 2025)", level=3)
add_p(
    "Em novembro de 2025, o Banco Central publicou três resoluções que estruturam o "
    "arcabouço regulatório das PSAVs:"
)

doc.add_heading("Resolução BCB nº 519 — Autorização e Funcionamento", level=4)
add_p(
    "Estabelece os requisitos para que uma PSAV obtenha autorização de funcionamento junto "
    "ao Banco Central. Inclui exigências de capital mínimo, governança corporativa, "
    "idoneidade dos controladores e administradores, e plano de negócios detalhado."
)

doc.add_heading("Resolução BCB nº 520 — Governança e Gestão de Riscos", level=4)
add_p(
    "Define regras de governança, gestão de riscos operacionais, cibernéticos e de liquidez, "
    "controles internos e auditoria para PSAVs. Alinha o setor cripto aos padrões de "
    "compliance já exigidos de instituições financeiras tradicionais."
)

doc.add_heading("Resolução BCB nº 521 — Segregação Patrimonial", level=4)
add_p(
    "Disciplina a obrigação de separar os recursos e criptoativos dos clientes do patrimônio "
    "próprio da PSAV. Nos termos da Resolução BCB nº 521, os ativos dos clientes não se "
    "comunicam com as obrigações da empresa, mesmo em caso de falência ou recuperação "
    "judicial da prestadora."
)

doc.add_heading("4.3 Resolução CMN 5.280 — Equiparação a Instituição Financeira", level=3)
add_p(
    "Em fevereiro de 2026, a Resolução CMN nº 5.280 deu um passo adicional: as PSAVs "
    "autorizadas pelo Banco Central passaram a ser equiparadas a instituições financeiras para "
    "fins da Lei Complementar nº 105/2001 (Lei do Sigilo Bancário)."
)
add_p("As implicações práticas dessa equiparação incluem:")
equip = [
    "O Banco Central pode acessar dados de clientes de PSAVs nas mesmas condições em que "
    "acessa dados de clientes de bancos.",
    "A Receita Federal pode requisitar informações de PSAVs mediante procedimento fiscal, "
    "nos termos do art. 6º da LC 105/2001.",
    "O sigilo bancário aplica-se integralmente às PSAVs — dados de clientes não podem ser "
    "divulgados a terceiros sem autorização legal."
]
for e in equip:
    add_bullet(e)

add_legal_box(
    "Resolução CMN nº 5.280, de fevereiro de 2026, c/c Lei Complementar nº 105/2001."
)

doc.add_heading("4.4 Quando a CVM Atua — E Quando Não Atua", level=3)
add_p(
    "A Comissão de Valores Mobiliários (CVM) NÃO é a reguladora padrão de criptoativos ou "
    "de PSAVs. A CVM somente tem competência quando o ativo virtual se enquadra como valor "
    "mobiliário (security), nos termos do art. 2º da Lei 6.385/1976.",
    bold=True
)

add_p("Na prática, a CVM atua em situações como:")
cvm_atua = [
    "Tokens que representam participação em empreendimentos (equity tokens).",
    "Tokens que oferecem direito a rendimentos futuros derivados de esforço de terceiros "
    "(security tokens).",
    "Ofertas públicas de tokens que se enquadrem no conceito de distribuição de valores "
    "mobiliários."
]
for c in cvm_atua:
    add_bullet(c)

add_p("A CVM NÃO regula, por padrão:")
cvm_nao = [
    "Bitcoin (BTC), Ethereum (ETH) e demais criptoativos de uso geral.",
    "Stablecoins (salvo se se enquadrarem como valor mobiliário).",
    "Exchanges e PSAVs como categoria — essa competência é do Banco Central."
]
for c in cvm_nao:
    add_bullet(c)

add_p(
    "É comum encontrar referências genéricas a \"registro na CVM\" como requisito para operar "
    "com criptoativos. Essa afirmação é imprecisa. O registro obrigatório para PSAVs é junto "
    "ao Banco Central, nos termos do Decreto 11.563/2023 e da Resolução BCB nº 519.",
    bold=True
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# CAPÍTULO 5
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("Capítulo 5", level=1)
doc.add_heading("Tributação de Criptoativos: Guia com Árvore Decisória", level=2)

add_p(
    "A tributação de criptoativos no Brasil depende de múltiplas variáveis: jurisdição da "
    "custódia, tipo de operação, valor alienado, e natureza do ativo. Este capítulo apresenta "
    "uma árvore decisória para que o investidor identifique o tratamento tributário aplicável "
    "a cada situação."
)

doc.add_heading("5.1 Árvore Decisória — Fluxo de Tributação", level=3)

add_p("Para determinar o tratamento tributário de uma operação com criptoativos, siga a seguinte sequência de perguntas:", bold=True)

doc.add_heading("PERGUNTA 1: Onde está a custódia?", level=4)
add_p("A) Custódia no Brasil (exchange brasileira ou self-custody com residência fiscal no Brasil)")
add_p("B) Custódia no exterior (exchange estrangeira ou self-custody em jurisdição estrangeira)")
add_p("→ Se B: vá para a seção 5.3 (Lei 14.754/2023).")
add_p("→ Se A: continue para a Pergunta 2.")

doc.add_heading("PERGUNTA 2: Qual o tipo de operação?", level=4)
add_p("A) Venda (alienação) de criptoativo por reais ou outra moeda fiduciária")
add_p("B) Permuta (troca de um criptoativo por outro)")
add_p("C) Staking, yield farming, liquidity mining")
add_p("D) Airdrop")
add_p("E) Doação ou herança")
add_p("→ Cada tipo tem tratamento específico. Veja seções abaixo.")

doc.add_heading("PERGUNTA 3 (para vendas em exchanges brasileiras): O valor total alienado no mês ultrapassou R$35 mil?", level=4)
add_p(
    "A isenção de R$35.000,00 mensais em alienações de criptoativos exige o cumprimento "
    "simultâneo de condições específicas. NÃO é uma isenção incondicional.",
    bold=True
)

doc.add_heading("5.2 Regra dos R$35 Mil — Condições e Limitações", level=3)

add_p(
    "A isenção para ganhos de capital em alienações de criptoativos de até R$35.000,00 mensais "
    "é frequentemente apresentada de forma simplificada. Na prática, conforme a Solução de "
    "Consulta Cosit nº 214/2021 e orientações da Receita Federal, ela se aplica sob condições "
    "específicas:"
)

doc.add_heading("Condições para a isenção:", level=4)
condicoes_35k = [
    "O valor de R$35.000,00 refere-se ao total de alienações no mês, não ao lucro. "
    "Se o investidor vender R$36.000,00 em criptoativos no mês, mesmo que o lucro tenha "
    "sido de R$100,00, a isenção NÃO se aplica.",
    "A isenção aplica-se a operações realizadas em exchanges brasileiras (domiciliadas no Brasil).",
    "Operações realizadas em exchanges estrangeiras seguem as regras da Lei 14.754/2023, "
    "que não prevê essa isenção — a tributação é de 15% sobre o rendimento, independentemente "
    "do valor.",
    "Self-custody: a análise depende da natureza da operação. Se o investidor opera via "
    "protocolo DeFi estrangeiro, a operação pode ser caracterizada como realizada no exterior.",
    "Permutas (troca de cripto por cripto) são, segundo a Receita Federal, fato gerador de "
    "ganho de capital. O valor de referência para fins do limite de R$35 mil é o valor "
    "de mercado do ativo recebido no momento da permuta."
]
for c in condicoes_35k:
    add_bullet(c)

add_legal_box(
    "Solução de Consulta Cosit nº 214/2021. IN RFB nº 1.888/2019 (vigente até jun/2026). "
    "Perguntão IRPF — Receita Federal (seção sobre criptoativos)."
)

doc.add_heading("Árvore decisória dos R$35 mil:", level=4)

tree = [
    "1. A operação foi uma alienação (venda ou permuta)? → Se NÃO, a regra não se aplica.",
    "2. A operação foi realizada em exchange brasileira? → Se NÃO (exterior), isenção NÃO se aplica — tributação pela Lei 14.754/2023.",
    "3. O total de alienações no mês foi ≤ R$35.000,00? → Se NÃO, tributação normal sobre o ganho de capital.",
    "4. Houve ganho de capital (preço de venda > custo de aquisição)? → Se NÃO, não há imposto (prejuízo).",
    "5. Se todas as condições acima forem atendidas → ganho de capital ISENTO."
]
for t in tree:
    add_p(t)

doc.add_heading("5.3 Custódia no Exterior — Lei 14.754/2023", level=3)

add_p(
    "A Lei 14.754/2023, regulamentada pela IN RFB 2.180/2024, alterou significativamente "
    "o tratamento tributário de aplicações financeiras no exterior por pessoas físicas "
    "residentes no Brasil."
)

add_p("Nos termos da Lei 14.754/2023:")
lei_14754 = [
    "Rendimentos de aplicações financeiras no exterior, incluindo criptoativos, são tributados "
    "à alíquota de 15%, apurados anualmente.",
    "A tributação incide sobre o rendimento (ganho), independentemente do valor de alienação "
    "— a isenção de R$35 mil NÃO se aplica a ativos custodiados no exterior.",
    "Aplicações em entidades controladas no exterior (offshore) seguem regras específicas de "
    "tributação na pessoa controladora.",
    "Os rendimentos devem ser declarados no ajuste anual do IRPF."
]
for l in lei_14754:
    add_bullet(l)

add_legal_box("Lei nº 14.754, de 12 de dezembro de 2023. IN RFB nº 2.180, de 2024.")

doc.add_heading("5.4 Permutas", level=3)
add_p(
    "De acordo com a orientação da Receita Federal, a permuta de criptoativos (troca de um "
    "ativo por outro) é considerada fato gerador de ganho de capital. O valor de mercado do "
    "ativo recebido no momento da permuta é utilizado como referência para apuração do ganho."
)
add_p(
    "Exemplo: se o investidor troca 1 BTC (custo de aquisição: R$200.000) por ETH no valor "
    "de mercado de R$250.000 no momento da permuta, há ganho de capital de R$50.000, sujeito "
    "à tributação."
)

doc.add_heading("5.5 Staking, Yield Farming e Airdrops", level=3)
add_p(
    "A Receita Federal ainda não publicou orientação definitiva consolidada sobre a tributação "
    "de todos os tipos de recompensas em criptoativos. No entanto, com base nas orientações "
    "disponíveis e na jurisprudência administrativa:"
)
staking_rules = [
    "Staking: recompensas de staking tendem a ser tratadas como rendimento tributável no "
    "momento do recebimento, pelo valor de mercado na data. O custo de aquisição para "
    "futuras alienações é o valor na data do recebimento.",
    "Yield farming e liquidity mining: tratamento análogo ao staking — rendimento tributável "
    "no recebimento.",
    "Airdrops: o custo de aquisição pode ser considerado zero. Na alienação futura, o ganho "
    "de capital seria o valor total da venda.",
    "DeFi: operações via protocolos DeFi internacionais podem ser enquadradas como operações "
    "no exterior, sujeitas à Lei 14.754/2023."
]
for s in staking_rules:
    add_bullet(s)

add_p(
    "Importante: dada a ausência de regulamentação específica para cada tipo de operação "
    "DeFi, recomenda-se consultar um tributarista especializado para operações complexas.",
    italic=True
)

doc.add_heading("5.6 Alíquotas de Ganho de Capital", level=3)
add_p("Para operações domésticas (em exchanges brasileiras), as alíquotas progressivas de ganho de capital são:")
aliquotas = [
    "Até R$5 milhões: 15%",
    "De R$5 milhões a R$10 milhões: 17,5%",
    "De R$10 milhões a R$30 milhões: 20%",
    "Acima de R$30 milhões: 22,5%"
]
for a in aliquotas:
    add_bullet(a)

add_p("Para operações no exterior (Lei 14.754/2023): alíquota fixa de 15%.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# CAPÍTULO 6
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("Capítulo 6", level=1)
doc.add_heading("Rearp: O Que Foi e O Que Significa Agora", level=2)

add_p(
    "O Regime Especial de Regularização Cambial e Tributária (Rearp), também conhecido como "
    "\"Repatriação de Criptoativos\", foi uma oportunidade de regularização que já se encerrou. "
    "Este capítulo trata o Rearp como fato histórico e analisa suas consequências para quem "
    "aderiu e para quem não aderiu.",
    bold=True
)

doc.add_heading("6.1 O Que Foi o Rearp", level=3)
add_p(
    "O Rearp foi instituído como parte da Lei 14.754/2023, em conjunto com dispositivos "
    "posteriores que regulamentaram a adesão. Tratava-se de um regime especial que permitia "
    "a pessoas físicas residentes no Brasil regularizar ativos mantidos no exterior — incluindo "
    "criptoativos — que não haviam sido declarados ou que apresentavam inconsistências "
    "na declaração."
)
add_p("As condições do Rearp incluíam:")
rearp_cond = [
    "Alíquota reduzida sobre o valor dos ativos regularizados.",
    "Possibilidade de declarar ativos anteriormente omitidos sem penalidades de sonegação.",
    "Adesão voluntária e irretratável."
]
for r in rearp_cond:
    add_bullet(r)

doc.add_heading("6.2 A Janela Já Se Encerrou", level=3)
add_p(
    "O prazo para adesão ao Derp (Declaração de Regularização) expirou em 19 de fevereiro "
    "de 2026, com pagamento devido até 27 de fevereiro de 2026. Não há previsão de "
    "reabertura até a data de publicação deste material (abril de 2026).",
    bold=True
)
add_p(
    "Isso significa que investidores que não aderiram ao Rearp durante a janela de "
    "regularização não podem mais utilizar esse mecanismo para regularizar ativos."
)

doc.add_heading("6.3 O Que Significa Para Quem Aderiu", level=3)
add_p("Investidores que aderiram ao Rearp dentro do prazo:")
aderiu = [
    "Regularizaram a situação dos ativos declarados, eliminando o risco de autuação por "
    "omissão referente aos períodos cobertos.",
    "Pagaram a alíquota reduzida sobre o valor dos ativos.",
    "Devem manter os ativos regularizados devidamente declarados nas próximas declarações "
    "de ajuste anual.",
    "A adesão não isenta de obrigações futuras — a DeCripto e demais obrigações acessórias "
    "continuam aplicáveis."
]
for a in aderiu:
    add_bullet(a)

doc.add_heading("6.4 O Que Significa Para Quem Não Aderiu", level=3)
add_p("Investidores que não aderiram ao Rearp devem considerar:")
nao_aderiu = [
    "A não adesão, por si só, não configura infração — o Rearp era voluntário.",
    "No entanto, se existem ativos não declarados ou com inconsistências, o risco de autuação "
    "permanece e tende a aumentar com a entrada em vigor da DeCripto (julho de 2026) e das "
    "trocas internacionais via CARF (2027).",
    "O investidor que possui ativos não regularizados deve buscar orientação jurídica "
    "especializada para avaliar as opções disponíveis dentro da legislação vigente.",
    "A regularização espontânea (denúncia espontânea), nos termos do art. 138 do CTN, pode "
    "ser uma alternativa — desde que realizada antes de qualquer procedimento fiscal "
    "relacionado."
]
for n in nao_aderiu:
    add_bullet(n)

add_legal_box(
    "Lei nº 14.754/2023 (dispositivos sobre regularização). "
    "Art. 138 do Código Tributário Nacional (denúncia espontânea)."
)

add_boost_opinion(
    "Para investidores que não aderiram ao Rearp e possuem ativos com situação irregular, "
    "o tempo para buscar orientação jurídica está se esgotando. Com a DeCripto entrando em "
    "vigor em julho de 2026, a capacidade de detecção de inconsistências pela Receita Federal "
    "aumentará significativamente."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# CAPÍTULO 7
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("Capítulo 7", level=1)
doc.add_heading("Cenários de Multa e Penalidades", level=2)

add_p(
    "Este capítulo apresenta os principais cenários de penalidade previstos na legislação "
    "vigente. Todas as informações estão condicionadas à interpretação e aplicação pelos "
    "órgãos competentes — as multas descritas são as previstas em lei, não determinações "
    "automáticas."
)

doc.add_heading("7.1 Omissão de Informações na DeCripto", level=3)
add_p(
    "De acordo com a IN RFB 2.291/2025, a omissão ou prestação de informações incorretas "
    "na DeCripto pode acarretar multa de até 3% do valor da operação omitida ou informada "
    "incorretamente."
)
add_p(
    "Exemplo hipotético: se um investidor omitiu operações de compra e venda no valor total "
    "de R$500.000,00, a multa pode chegar a R$15.000,00 (3% × R$500.000), nos termos "
    "da IN RFB 2.291/2025."
)
add_legal_box("IN RFB nº 2.291/2025 — disposições sobre penalidades.")

doc.add_heading("7.2 Não Recolhimento de Imposto sobre Ganho de Capital", level=3)
add_p(
    "O ganho de capital apurado em operações com criptoativos deve ser recolhido via DARF "
    "até o último dia útil do mês subsequente à operação (para operações domésticas). "
    "O não recolhimento no prazo pode acarretar, nos termos da legislação tributária:"
)
multas_gc = [
    "Multa de mora: 0,33% ao dia, limitada a 20% do valor do imposto.",
    "Juros de mora: calculados pela taxa Selic acumulada.",
    "Em caso de lançamento de ofício pela Receita Federal: multa de 75% do valor do imposto "
    "devido (art. 44 da Lei 9.430/1996).",
    "Em caso de fraude, sonegação ou conluio: a multa pode chegar a 150% do valor do imposto "
    "(art. 44, §1º da Lei 9.430/1996)."
]
for m in multas_gc:
    add_bullet(m)

add_legal_box("Art. 44 da Lei nº 9.430/1996. CTN, arts. 136 a 138.")

doc.add_heading("7.3 Omissão na Declaração de Ajuste Anual (DIRPF)", level=3)
add_p(
    "Criptoativos devem ser declarados na ficha de Bens e Direitos da DIRPF quando o valor "
    "de aquisição de cada tipo de criptoativo exceder os limites definidos pela Receita "
    "Federal. A omissão pode acarretar:"
)
multas_dirpf = [
    "Malha fina e intimação para esclarecimentos.",
    "Multa por omissão de rendimentos: 75% sobre a diferença de imposto apurada (art. 44 "
    "da Lei 9.430/1996).",
    "Em casos de dolo: 150% (agravamento)."
]
for m in multas_dirpf:
    add_bullet(m)

doc.add_heading("7.4 Operações no Exterior Não Declaradas", level=3)
add_p(
    "Nos termos da Lei 14.754/2023, rendimentos de aplicações financeiras no exterior "
    "— incluindo criptoativos — que não forem declarados podem acarretar, além das "
    "penalidades tributárias acima, enquadramento em:"
)
exterior_pen = [
    "Evasão de divisas (Lei 7.492/1986, art. 22): quando há manutenção de valores no "
    "exterior sem a devida declaração. Pena: reclusão de 2 a 6 anos, e multa.",
    "Sonegação fiscal (Lei 8.137/1990): quando há omissão dolosa de rendimentos tributáveis. "
    "Pena: reclusão de 2 a 5 anos, e multa."
]
for e in exterior_pen:
    add_bullet(e)

add_p(
    "Nota: as penalidades penais acima se aplicam apenas em casos de dolo comprovado "
    "(intenção de omitir). Erros de boa-fé, quando corrigidos espontaneamente, não "
    "configuram crime tributário.",
    italic=True
)

doc.add_heading("7.5 Linguagem Condicional — Aviso Importante", level=3)
add_p(
    "Todos os cenários descritos neste capítulo utilizam linguagem condicional (\"pode "
    "acarretar\", \"nos termos de\") porque a aplicação de penalidades depende de "
    "procedimento fiscal específico, com direito a contraditório e ampla defesa. "
    "Nenhuma multa é automática — todas dependem de lançamento pela autoridade "
    "competente e podem ser contestadas administrativamente.",
    bold=True
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# CAPÍTULO 8
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("Capítulo 8", level=1)
doc.add_heading("Checklist: 7 Passos para Conformidade", level=2)

add_p(
    "Este checklist prático resume as ações recomendadas para que o investidor com patrimônio "
    "em criptoativos esteja em conformidade antes da entrada em vigor da DeCripto em julho "
    "de 2026."
)

passos = [
    ("Passo 1: Levante seu histórico completo de operações",
     "Prazo recomendado: até maio de 2026",
     [
         "Consolide todas as operações realizadas desde o início dos investimentos em criptoativos.",
         "Inclua: compras, vendas, permutas, staking, airdrops, transferências entre wallets.",
         "Identifique as exchanges utilizadas (brasileiras e estrangeiras) e os períodos de operação.",
         "Organize por ano fiscal para facilitar a apuração de ganhos de capital."
     ]),
    ("Passo 2: Identifique a jurisdição de custódia de cada ativo",
     "Prazo recomendado: até maio de 2026",
     [
         "Classifique cada ativo como: custódia Brasil, custódia exterior ou self-custody.",
         "Para self-custody, avalie se as operações foram realizadas via protocolos domésticos ou estrangeiros.",
         "A jurisdição de custódia determina o regime tributário aplicável (ver Capítulo 5)."
     ]),
    ("Passo 3: Apure ganhos de capital pendentes",
     "Prazo recomendado: até maio de 2026",
     [
         "Calcule o ganho de capital de cada operação tributável.",
         "Verifique se a isenção de R$35 mil se aplica a cada caso (ver árvore decisória, Capítulo 5).",
         "Identifique DARFs não recolhidas e juros/multa de mora acumulados.",
         "Para operações no exterior: apure os rendimentos nos termos da Lei 14.754/2023."
     ]),
    ("Passo 4: Regularize pendências tributárias",
     "Prazo recomendado: até junho de 2026",
     [
         "Recolha DARFs pendentes com os acréscimos legais (multa de mora + Selic).",
         "Se houver omissão de rendimentos em declarações anteriores, considere retificar as DIRPFs.",
         "A denúncia espontânea (art. 138, CTN) pode ser utilizada antes de procedimento fiscal.",
         "Consulte um tributarista para operações complexas ou valores relevantes."
     ]),
    ("Passo 5: Atualize a Declaração de Bens e Direitos",
     "Prazo recomendado: até a entrega da DIRPF 2026 (referente ao ano-calendário 2025)",
     [
         "Declare todos os criptoativos na ficha de Bens e Direitos, usando os códigos corretos.",
         "Informe o custo de aquisição de cada ativo (não o valor de mercado).",
         "Inclua ativos em self-custody e em exchanges estrangeiras.",
         "Preencha a Declaração de Capitais Brasileiros no Exterior (CBE) se aplicável — obrigatória "
         "para ativos no exterior acima de USD 1 milhão (Banco Central)."
     ]),
    ("Passo 6: Prepare-se para a DeCripto",
     "Prazo: até 1º de julho de 2026",
     [
         "Certifique-se de que suas exchanges brasileiras possuem seus dados cadastrais atualizados.",
         "Se opera em exchanges estrangeiras: prepare-se para reportar operações via DeCripto "
         "quando o valor mensal ultrapassar o limite estabelecido pela IN RFB 2.291/2025.",
         "Organize um sistema de registro contínuo de operações para facilitar o reporte mensal.",
         "Considere utilizar ferramentas especializadas de tracking de operações cripto."
     ]),
    ("Passo 7: Monitore desenvolvimentos regulatórios",
     "Prazo: contínuo",
     [
         "Acompanhe regulamentações complementares do Banco Central sobre PSAVs.",
         "Monitore a adesão de novos países ao CARF para avaliar riscos de troca de informações.",
         "Verifique atualizações da Receita Federal sobre a DeCripto e eventuais ajustes nos prazos.",
         "Fique atento a decisões do CARF (Conselho Administrativo de Recursos Fiscais) sobre "
         "tributação de criptoativos — a jurisprudência administrativa está em formação."
     ])
]

for titulo, prazo, items in passos:
    doc.add_heading(titulo, level=3)
    add_p(prazo, bold=True, italic=True, color=(0x0B, 0x6E, 0x4F))
    for item in items:
        add_bullet(item)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# CAPÍTULO 9
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("Capítulo 9", level=1)
doc.add_heading("Próximos Passos", level=2)

add_p(
    "Esta seção apresenta a visão da Boost Research sobre como o investidor pode se "
    "posicionar diante do novo cenário regulatório. O conteúdo abaixo reflete a opinião "
    "da empresa e tem caráter comercial.",
    italic=True, color=(0x99, 0x99, 0x99)
)

doc.add_heading("9.1 O Cenário Pede Ação, Não Reação", level=3)
add_boost_opinion(
    "O investidor que aguarda a DeCripto entrar em vigor para se organizar corre o risco "
    "de enfrentar um acúmulo de pendências — histórico de operações sem registro, ganhos "
    "de capital não apurados, ativos não declarados. Na Boost Research, acreditamos que a "
    "antecipação é a estratégia mais racional."
)

doc.add_heading("9.2 O Que a Mentoria Boost Research Oferece", level=3)
add_p(
    "A Mentoria Boost Research é um programa de acompanhamento para investidores de "
    "criptoativos que desejam tomar decisões mais informadas sobre seu portfólio."
)
add_p("Na Mentoria Boost Research, nossos membros têm acesso a:", bold=True)
mentoria = [
    "Análises semanais sobre o mercado de criptoativos, com foco em fundamentos e dados on-chain.",
    "Portfólio sugerido com teses de investimento fundamentadas e atualizadas.",
    "Acompanhamento regulatório: monitoramos as mudanças na legislação e explicamos o impacto "
    "prático para o investidor.",
    "Comunidade de investidores qualificados para troca de experiências e networking.",
    "Acesso direto a André Franco e à equipe de analistas para dúvidas sobre posicionamento "
    "de portfólio."
]
for m in mentoria:
    add_bullet(m)

add_p(
    "Importante: a Mentoria Boost Research NÃO oferece aconselhamento jurídico ou tributário "
    "individualizado. Para questões legais específicas, recomendamos a consulta a profissionais "
    "habilitados (advogados tributaristas e contadores especializados).",
    bold=True
)

doc.add_heading("9.3 Como Saber Mais", level=3)
add_boost_opinion(
    "Se você deseja aprofundar sua análise sobre o mercado de criptoativos e acompanhar as "
    "mudanças regulatórias com orientação de especialistas, conheça a Mentoria Boost Research. "
    "Acesse: boostresearch.com.br"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# REFERÊNCIAS E FONTES
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("Referências e Fontes", level=1)

doc.add_heading("Legislação Federal", level=3)
leis = [
    "Lei nº 6.385, de 7 de dezembro de 1976 — Dispõe sobre o mercado de valores mobiliários e cria a CVM.",
    "Lei Complementar nº 105, de 10 de janeiro de 2001 — Sigilo das operações de instituições financeiras.",
    "Medida Provisória nº 2.158-35, de 24 de agosto de 2001 — Obrigações acessórias e penalidades.",
    "Lei nº 7.492, de 16 de junho de 1986 — Crimes contra o sistema financeiro nacional.",
    "Lei nº 8.137, de 27 de dezembro de 1990 — Crimes contra a ordem tributária.",
    "Lei nº 9.430, de 27 de dezembro de 1996 — Legislação tributária federal (art. 44: multas de ofício).",
    "Código Tributário Nacional (Lei nº 5.172/1966) — Arts. 136 a 138 (denúncia espontânea).",
    "Lei nº 14.478, de 21 de dezembro de 2022 — Marco Legal dos Ativos Virtuais.",
    "Lei nº 14.754, de 12 de dezembro de 2023 — Tributação de aplicações financeiras no exterior."
]
for l in leis:
    add_bullet(l)

doc.add_heading("Decretos", level=3)
decretos = [
    "Decreto nº 11.563, de 13 de junho de 2023 — Designa o Banco Central como regulador das PSAVs."
]
for d in decretos:
    add_bullet(d)

doc.add_heading("Instruções Normativas da Receita Federal", level=3)
ins = [
    "IN RFB nº 1.888, de 3 de maio de 2019 — Obrigação de prestação de informações sobre operações "
    "com criptoativos (vigente até junho de 2026).",
    "IN RFB nº 2.180, de 2024 — Regulamenta a Lei 14.754/2023 (aplicações no exterior).",
    "IN RFB nº 2.291, de 22 de janeiro de 2025 — Institui a Declaração de Criptoativos (DeCripto)."
]
for i in ins:
    add_bullet(i)

doc.add_heading("Resoluções do Banco Central e CMN", level=3)
resolucoes = [
    "Resolução BCB nº 519, de novembro de 2025 — Autorização e funcionamento de PSAVs.",
    "Resolução BCB nº 520, de novembro de 2025 — Governança e gestão de riscos de PSAVs.",
    "Resolução BCB nº 521, de novembro de 2025 — Segregação patrimonial de PSAVs.",
    "Resolução CMN nº 5.280, de fevereiro de 2026 — PSAVs como instituições financeiras sob LC 105."
]
for r in resolucoes:
    add_bullet(r)

doc.add_heading("Soluções de Consulta e Orientações", level=3)
solucoes = [
    "Solução de Consulta Cosit nº 214/2021 — Tratamento tributário de criptoativos.",
    "Perguntão IRPF — Receita Federal do Brasil (seção sobre criptoativos, atualizado anualmente)."
]
for s in solucoes:
    add_bullet(s)

doc.add_heading("Fontes Internacionais", level=3)
internacionais = [
    "OCDE — Crypto-Asset Reporting Framework (CARF), 2023. Disponível em: oecd.org/tax/automatic-exchange",
    "OCDE — Common Reporting Standard (CRS). Disponível em: oecd.org/tax/automatic-exchange/common-reporting-standard",
    "G20 — Declaração de compromisso com implementação do CARF (2023)."
]
for i in internacionais:
    add_bullet(i)

doc.add_heading("Portais Oficiais", level=3)
portais = [
    "Receita Federal do Brasil: gov.br/receitafederal",
    "Banco Central do Brasil: bcb.gov.br",
    "Comissão de Valores Mobiliários: gov.br/cvm",
    "OCDE — Automatic Exchange Portal: oecd.org/tax/automatic-exchange"
]
for p in portais:
    add_bullet(p)

# ── Final ──
doc.add_paragraph()
doc.add_paragraph()
add_p("—", align=WD_ALIGN_PARAGRAPH.CENTER)
add_p(
    "Este material foi elaborado por André Franco, fundador e analista da Boost Research, "
    "com base na legislação vigente até abril de 2026. Para orientação jurídica ou tributária "
    "individualizada, consulte um profissional habilitado.",
    size=10, italic=True, color=(0x99, 0x99, 0x99), align=WD_ALIGN_PARAGRAPH.CENTER
)
add_p(
    "© 2026 Boost Research. Todos os direitos reservados.",
    size=9, color=(0x99, 0x99, 0x99), align=WD_ALIGN_PARAGRAPH.CENTER
)

# ── Save ──
doc.save(OUTPUT_PATH)
print(f"E-book salvo em: {OUTPUT_PATH}")
